#!/usr/bin/env python3
import sys
import re
import subprocess
import requests
from pathlib import Path
from docx import Document
from docx.shared import Inches

def download(url, dest):
    r = requests.get(url, timeout=20)
    r.raise_for_status()
    with open(dest, "wb") as f:
        f.write(r.content)

def convert_mermaid(md_path):
    content = md_path.read_text()
    pattern = re.compile(r"```mermaid\s*(.*?)```", re.DOTALL)
    matches = pattern.findall(content)
    if not matches:
        return

    assets_dir = md_path.parent / (md_path.stem + "_assets")
    assets_dir.mkdir(exist_ok=True)

    for i, code in enumerate(matches, start=1):
        mmd = assets_dir / f"diagram_{i}.mmd"
        png = assets_dir / f"diagram_{i}.png"
        mmd.write_text(code.strip())
        subprocess.run(["mmdc", "-i", str(mmd), "-o", str(png)], check=True)
        content = content.replace(f"```mermaid\n{code}```", f"![Diagram]({png.name})")

    md_path.write_text(content)

def md_to_docx(md_file, docx_file, sap_url, mm_url):
    convert_mermaid(md_file)

    sap_file = md_file.parent / "sap_logo.png"
    mm_file = md_file.parent / "mm_logo.png"
    download(sap_url, sap_file)
    download(mm_url, mm_file)

    subprocess.run(["pandoc", str(md_file), "-o", str(docx_file)], check=True)

    doc = Document(str(docx_file))
    header = doc.sections[0].header
    table = header.add_table(rows=1, cols=2)

    table.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(sap_file), width=Inches(1.5))
    p = table.rows[0].cells[1].paragraphs[0]
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    except:
        pass
    p.add_run().add_picture(str(mm_file), width=Inches(1.5))

    doc.save(str(docx_file))

if __name__ == "__main__":
    md_to_docx(Path(sys.argv[1]), Path(sys.argv[2]), sys.argv[3], sys.argv[4])

