import os
import json
import sys
import xml.etree.ElementTree as ET 
from openai import OpenAI
from docx import Document # <<< ADDED IMPORT for DOCX manipulation

# --- Configuration ---
INPUT_PACKAGE_NAME = os.environ.get("INPUT_PACKAGE_NAME")
INPUT_IFLOW_NAME = os.environ.get("INPUT_IFLOW_NAME")
GITHUB_WORKSPACE_ROOT = os.environ.get("GITHUB_WORKSPACE_ROOT") 
openai_api_key = os.environ.get("OPENAI_API_KEY")

# --- PATHS ---
DOCX_TEMPLATE_PATH = '../../assets/project_template.docx' 

# Define output path components:
RELATIVE_OUTPUT_PATH_DIR = f'cpi-artifacts/{INPUT_PACKAGE_NAME}/{INPUT_IFLOW_NAME}'
# *** CRITICAL: Output is now DOCX to retain styling ***
OUTPUT_FILENAME_DOCX = f'{INPUT_IFLOW_NAME}_Technical_Spec.docx' 

# Path to the source .iflw file:
IFLOW_SOURCE_FILENAME = f'{INPUT_IFLOW_NAME}.iflw' 
NESTED_IFLOW_SUBPATH = f'src/main/resources/scenarioflows/integrationflow/{IFLOW_SOURCE_FILENAME}'
RELATIVE_IFLOW_SOURCE_PATH = os.path.join(RELATIVE_OUTPUT_PATH_DIR, NESTED_IFLOW_SUBPATH)

# Project-specific context
PROJECT_CONTEXT = {
    "iflow_name": INPUT_IFLOW_NAME,
    "package_name": INPUT_PACKAGE_NAME,
    "author": "Rohancherian783",
    "source_system": "SAP S/4HANA (IDoc)",
    "target_system": "Internal HR System (REST API)",
    "integration_scenario": "Asynchronous Employee Data Replication",
    "security_protocol": "Basic Authentication over HTTPS"
}

# --- Helper Functions ---

# NOTE: extract_text_from_docx is now just used to get template structure for AI's context
def extract_text_from_docx(file_path):
    """Reads a .docx file and returns its content as a single string."""
    try:
        if not os.path.exists(file_path):
            print(f"ERROR: Word template not found at {file_path}")
            sys.exit(1)
        document = Document(file_path)
        full_text = [p.text for p in document.paragraphs if p.text.strip()]
        return '\n\n'.join(full_text)
    except Exception as e:
        print(f"ERROR: Failed to read DOCX file: {e}")
        sys.exit(1)

def call_openai_api(prompt, api_key):
    """Calls the OpenAI API to generate the document content in MARKDOWN."""
    try:
        client = OpenAI(api_key=api_key)
        model = "gpt-4o-mini" 

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": 
                 "You are an expert CPI Technical Writer. Your task is to generate the document body "
                 "in **CLEAN MARKDOWN** format. Use '#' for main sections, '##' for subsections, and lists where appropriate. "
                 "**DO NOT** use any HTML tags or attempt to replicate DOCX styling. Your output will be text/markdown only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
        )
        ai_generated_content = response.choices[0].message.content
        return ai_generated_content
        
    except Exception as e:
        print(f"FATAL ERROR: OpenAI API call failed: {e}")
        return ""

def parse_iflow_xml(xml_content):
    """
    Parses the iFlow XML content to extract key technical details for the AI.
    *** CRITICAL: You MUST inspect your .iflw file and adjust these XPaths. ***
    """
    try:
        root = ET.fromstring(xml_content)
        details = {}
        
        # --- Example XPaths (ADJUST THESE!) ---
        receiver_address = root.find(".//endpoint/address") 
        details['Target_Endpoint_URL'] = receiver_address.text if receiver_address is not None and receiver_address.text else "N/A"
        
        sender_type = root.find(".//sender/messageProtocol")
        details['Sender_Protocol'] = sender_type.text if sender_type is not None else "HTTPS/IDoc"

        credential = root.find(".//security/credentialName")
        details['Security_Credential_Name'] = credential.text if credential is not None else "BasicAuth_Default"
        
        scripts = [s.get('name') for s in root.findall(".//script/resource")]
        details['Scripts_In_iFlow'] = ", ".join(scripts) if scripts else "None"

        mapping = root.find(".//messageMapping/resource")
        details['Mapping_Resource'] = mapping.get('name') if mapping is not None else "None (Groovy/XSLT only)"

        # Return a simple, structured Markdown/Text format
        return "\n".join([f"- **{k}:** {v}" for k, v in details.items()])

    except ET.ParseError as e:
        print(f"ERROR: Failed to parse iFlow XML. Check if the file is valid XML: {e}")
        return f"ERROR: Could not parse iFlow XML. Detail: {e}"
    except Exception as e:
        print(f"ERROR during custom XML analysis: {e}")
        return f"ERROR: Failed during detail extraction. Detail: {e}"


def read_iflow_source_data(file_path):
    """Reads the raw iFlow XML content and returns it for parsing."""
    absolute_path = os.path.join(GITHUB_WORKSPACE_ROOT, file_path)
    try:
        with open(absolute_path, 'r', encoding='utf-8') as f:
            return f.read()
            
    except Exception as e:
        print(f"FATAL ERROR: Could not read iFlow XML file at {absolute_path}: {e}")
        sys.exit(1)

def replace_placeholder_in_docx(doc, placeholder, new_text):
    """
    Finds a placeholder text string in the document and replaces it 
    with the new text, retaining the original paragraph style.
    
    NOTE: This is a simplified replacement. Multi-paragraph replacement 
    (like a full section) is much more complex and usually requires 
    replacing the entire paragraph with a run of new paragraphs.
    """
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, new_text)
            return True
    return False


# --- Main function ---
def main():
    
    # 1. Check for required environment variables
    if not all([INPUT_PACKAGE_NAME, INPUT_IFLOW_NAME, GITHUB_WORKSPACE_ROOT, openai_api_key]):
        print("FATAL ERROR: Required environment variables are missing.")
        sys.exit(1)

    # 2. Read Template Content and Raw XML
    template_context_text = extract_text_from_docx(DOCX_TEMPLATE_PATH)
    raw_iflow_xml = read_iflow_source_data(RELATIVE_IFLOW_SOURCE_PATH)

    # 3. Parse the XML to get a clean summary
    INTEGRATION_SOURCE_DATA = parse_iflow_xml(raw_iflow_xml)

    # 4. Construct the Prompt (Asking for MARKDOWN)
    context_str = json.dumps(PROJECT_CONTEXT, indent=2)
    prompt = (
        "**CRITICAL INSTRUCTIONS:**\n"
        "1. **PRIORITIZE** the IFLOW ANALYZED DATA.\n"
        "2. Generate the full document body in **CLEAN MARKDOWN** format, starting from the **1. Introduction**.\n"
        "3. Ensure the markdown structure (e.g., #, ##, lists) is clean and logical.\n"
        
        f"--- HIGH-LEVEL CONTEXT ---\n{context_str}\n\n"
        
        f"--- IFLOW ANALYZED DATA ---\n{INTEGRATION_SOURCE_DATA}\n\n"
        
        f"--- DOCUMENT TEMPLATE STRUCTURE (For content guidance only) ---\n{template_context_text}\n"
    )
    
    # 5. Generate Document Body (Receives Markdown Text)
    document_body_markdown = call_openai_api(prompt, openai_api_key)
    
    if not document_body_markdown.strip():
        print("ERROR: AI returned empty content. Failing the workflow.")
        sys.exit(1) 

    # 6. Assemble and Save the Final Document (DOCX manipulation)
    ABSOLUTE_OUTPUT_DIR = os.path.join(GITHUB_WORKSPACE_ROOT, RELATIVE_OUTPUT_PATH_DIR)
    ABSOLUTE_OUTPUT_FILE_PATH = os.path.join(ABSOLUTE_OUTPUT_DIR, OUTPUT_FILENAME_DOCX)
    
    os.makedirs(ABSOLUTE_OUTPUT_DIR, exist_ok=True)
    
    try:
        # Load the original template document
        document = Document(DOCX_TEMPLATE_PATH)
        
        # --- Simplified Template Population ---
        # NOTE: For full automation, you would need a more complex system to 
        # map Markdown sections to specific DOCX placeholders.
        # For this final version, we will replace the whole body section (after TOC/Cover).
        
        # Clear existing content after the first few paragraphs (assuming cover/TOC are first)
        for i in range(len(document.paragraphs) - 1, 3, -1):
            p = document.paragraphs[i]
            if not p.text.strip().startswith(('1.', '2.', '3.', '4.', '5.')):
                p._element.getparent().remove(p._element)


        # Append the AI-generated Markdown, letting python-docx handle the formatting
        # by creating new styled paragraphs based on the markdown headers.
        
        # For simplicity, we just add the Markdown content as new paragraphs
        # You will need to manually adjust the styles in the output DOCX if this is not perfect.
        for line in document_body_markdown.split('\n'):
            line = line.strip()
            if line.startswith('# '):
                # Apply Heading 1 style
                document.add_heading(line.lstrip('# ').strip(), level=1)
            elif line.startswith('## '):
                # Apply Heading 2 style
                document.add_heading(line.lstrip('## ').strip(), level=2)
            elif line.startswith('### '):
                # Apply Heading 3 style
                document.add_heading(line.lstrip('### ').strip(), level=3)
            elif line:
                # Add as normal text
                document.add_paragraph(line)

        document.save(ABSOLUTE_OUTPUT_FILE_PATH)
        
        print(f"✅ Successfully generated and saved styled DOCX document to {ABSOLUTE_OUTPUT_FILE_PATH}")

    except Exception as e:
        print(f"ERROR saving DOCX file: {e}")
        sys.exit(1)


    # 7. Save the relative path for the Bash script (Now using DOCX filename)
    RELATIVE_PATH_TO_SAVE = os.path.join(RELATIVE_OUTPUT_PATH_DIR, OUTPUT_FILENAME_DOCX)
    with open('output_path.txt', 'w') as f:
        f.write(RELATIVE_PATH_TO_SAVE)
    
    print(f"✅ Saved relative path to output_path.txt: {RELATIVE_PATH_TO_SAVE}")


if __name__ == "__main__":
    main()
